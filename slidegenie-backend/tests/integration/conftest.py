"""
Integration Test Configuration and Shared Fixtures.

This module provides shared fixtures and configuration for integration tests,
including database setup, test clients, and common test data.
"""
import asyncio
import os
import tempfile
from pathlib import Path
from typing import AsyncGenerator, Generator
from uuid import uuid4

import pytest
import pytest_asyncio
from httpx import AsyncClient
from sqlalchemy import text
from sqlalchemy.ext.asyncio import (
    AsyncEngine,
    AsyncSession,
    async_sessionmaker,
    create_async_engine,
)
from testcontainers.postgres import PostgresContainer
from testcontainers.redis import RedisContainer
from testcontainers.minio import MinioContainer

from app.core.config import Settings, get_settings
from app.infrastructure.database.base import Base
from app.main import app
from app.infrastructure.database.utils import get_db


# Test database name
TEST_DB_NAME = "slidegenie_test"


class TestSettings(Settings):
    """Test-specific settings override."""
    
    ENVIRONMENT: str = "test"
    DEBUG: bool = True
    
    # Override with test database
    POSTGRES_DB: str = TEST_DB_NAME
    
    # Disable rate limiting for tests
    RATE_LIMIT_ENABLED: bool = False
    
    # Use test email backend
    EMAIL_BACKEND: str = "test"
    
    # Shorter token expiry for testing
    ACCESS_TOKEN_EXPIRE_MINUTES: int = 5
    
    # Test file upload settings
    MAX_UPLOAD_SIZE: int = 10 * 1024 * 1024  # 10MB
    ALLOWED_UPLOAD_EXTENSIONS: list = [".pdf", ".docx", ".tex", ".latex"]
    
    # Test export settings
    EXPORT_QUEUE_ENABLED: bool = True
    EXPORT_CLEANUP_HOURS: int = 1


@pytest.fixture(scope="session")
def event_loop():
    """Create an instance of the default event loop for the test session."""
    loop = asyncio.get_event_loop_policy().new_event_loop()
    yield loop
    loop.close()


@pytest.fixture(scope="session")
def postgres_container():
    """Start PostgreSQL container for integration tests."""
    with PostgresContainer(
        image="pgvector/pgvector:pg16",
        username="test_user",
        password="test_password",
        dbname=TEST_DB_NAME,
    ) as postgres:
        # Initialize pgvector extension
        import psycopg2
        conn = psycopg2.connect(postgres.get_connection_url())
        with conn.cursor() as cursor:
            cursor.execute("CREATE EXTENSION IF NOT EXISTS vector;")
            cursor.execute("CREATE EXTENSION IF NOT EXISTS pg_trgm;")
        conn.commit()
        conn.close()
        
        yield postgres


@pytest.fixture(scope="session")
def redis_container():
    """Start Redis container for integration tests."""
    with RedisContainer(image="redis:7-alpine") as redis:
        yield redis


@pytest.fixture(scope="session")
def minio_container():
    """Start MinIO container for integration tests."""
    with MinioContainer(
        image="minio/minio:latest",
        access_key="test_access_key",
        secret_key="test_secret_key",
    ) as minio:
        # Create test bucket
        import boto3
        s3_client = boto3.client(
            "s3",
            endpoint_url=minio.get_connection_url(),
            aws_access_key_id=minio.access_key,
            aws_secret_access_key=minio.secret_key,
        )
        s3_client.create_bucket(Bucket="slidegenie-test")
        
        yield minio


@pytest.fixture(scope="session")
def test_settings(
    postgres_container,
    redis_container,
    minio_container,
) -> TestSettings:
    """Create test settings with container URLs."""
    # Parse PostgreSQL URL
    pg_url = postgres_container.get_connection_url()
    pg_url = pg_url.replace("postgresql://", "postgresql+asyncpg://")
    
    # Parse Redis URL
    redis_host, redis_port = redis_container.get_container_host_ip(), redis_container.get_exposed_port(6379)
    
    # Parse MinIO URL
    minio_host, minio_port = minio_container.get_container_host_ip(), minio_container.get_exposed_port(9000)
    
    return TestSettings(
        SECRET_KEY="test-secret-key-for-integration-testing-only",
        DATABASE_URL=pg_url,
        POSTGRES_USER=postgres_container.username,
        POSTGRES_PASSWORD=postgres_container.password,
        POSTGRES_HOST=postgres_container.get_container_host_ip(),
        POSTGRES_PORT=postgres_container.get_exposed_port(5432),
        REDIS_HOST=redis_host,
        REDIS_PORT=redis_port,
        MINIO_ENDPOINT=f"{minio_host}:{minio_port}",
        MINIO_ACCESS_KEY=minio_container.access_key,
        MINIO_SECRET_KEY=minio_container.secret_key,
        MINIO_BUCKET_NAME="slidegenie-test",
        ANTHROPIC_API_KEY="test-anthropic-key",
        OPENAI_API_KEY="test-openai-key",
        GOOGLE_CLIENT_ID="test-google-client-id",
        GOOGLE_CLIENT_SECRET="test-google-client-secret",
        MICROSOFT_CLIENT_ID="test-microsoft-client-id",
        MICROSOFT_CLIENT_SECRET="test-microsoft-client-secret",
    )


@pytest_asyncio.fixture(scope="session")
async def test_engine(test_settings: TestSettings) -> AsyncGenerator[AsyncEngine, None]:
    """Create test database engine."""
    engine = create_async_engine(
        str(test_settings.DATABASE_URL),
        echo=False,
        pool_pre_ping=True,
        pool_size=5,
        max_overflow=10,
    )
    
    # Create all tables
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)
    
    yield engine
    
    # Drop all tables
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.drop_all)
    
    await engine.dispose()


@pytest_asyncio.fixture
async def db_session(test_engine: AsyncEngine) -> AsyncGenerator[AsyncSession, None]:
    """Create database session for tests."""
    async_session = async_sessionmaker(
        test_engine,
        class_=AsyncSession,
        expire_on_commit=False,
    )
    
    async with async_session() as session:
        yield session
        await session.rollback()


@pytest.fixture
def override_settings(test_settings: TestSettings):
    """Override application settings with test settings."""
    app.dependency_overrides[get_settings] = lambda: test_settings
    yield
    app.dependency_overrides.clear()


@pytest_asyncio.fixture
async def override_db(db_session: AsyncSession):
    """Override database dependency with test session."""
    async def _get_test_db():
        yield db_session
    
    app.dependency_overrides[get_db] = _get_test_db
    yield
    app.dependency_overrides.pop(get_db, None)


@pytest_asyncio.fixture
async def client(
    override_settings,
    override_db,
) -> AsyncGenerator[AsyncClient, None]:
    """Create async HTTP client for testing."""
    async with AsyncClient(
        app=app,
        base_url="http://test",
        follow_redirects=False,
    ) as client:
        yield client


@pytest.fixture
def temp_upload_dir() -> Generator[Path, None, None]:
    """Create temporary directory for file uploads."""
    with tempfile.TemporaryDirectory() as tmpdir:
        yield Path(tmpdir)


@pytest.fixture
def sample_pdf_file(temp_upload_dir: Path) -> Path:
    """Create sample PDF file for testing."""
    file_path = temp_upload_dir / "test_presentation.pdf"
    
    # Create a minimal PDF
    pdf_content = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /Resources <<
/Font << /F1 4 0 R >>
>> /MediaBox [0 0 612 792] /Contents 5 0 R >>
endobj
4 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj
5 0 obj
<< /Length 44 >>
stream
BT
/F1 12 Tf
100 700 Td
(Test PDF Document) Tj
ET
endstream
endobj
xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000262 00000 n 
0000000341 00000 n 
trailer
<< /Size 6 /Root 1 0 R >>
startxref
435
%%EOF"""
    
    file_path.write_bytes(pdf_content)
    return file_path


@pytest.fixture
def sample_docx_file(temp_upload_dir: Path) -> Path:
    """Create sample DOCX file for testing."""
    from docx import Document
    
    file_path = temp_upload_dir / "test_presentation.docx"
    
    doc = Document()
    doc.add_heading("Test Presentation", 0)
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is a test document for SlideGenie integration testing.")
    doc.add_heading("Methods", level=1)
    doc.add_paragraph("We used automated testing to verify functionality.")
    doc.add_heading("Results", level=1)
    doc.add_paragraph("All tests passed successfully.")
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph("SlideGenie works as expected.")
    
    doc.save(file_path)
    return file_path


@pytest.fixture
def sample_latex_file(temp_upload_dir: Path) -> Path:
    """Create sample LaTeX file for testing."""
    file_path = temp_upload_dir / "test_presentation.tex"
    
    latex_content = r"""
\documentclass{beamer}
\usetheme{Madrid}
\title{Test Presentation}
\author{Test Author}
\institute{Test University}
\date{\today}

\begin{document}

\frame{\titlepage}

\begin{frame}
\frametitle{Introduction}
This is a test LaTeX presentation for integration testing.
\end{frame}

\begin{frame}
\frametitle{Methods}
\begin{itemize}
\item Automated testing
\item Integration verification
\item Performance monitoring
\end{itemize}
\end{frame}

\begin{frame}
\frametitle{Results}
\begin{equation}
E = mc^2
\end{equation}
All tests passed successfully.
\end{frame}

\begin{frame}
\frametitle{Conclusion}
SlideGenie LaTeX processing works correctly.
\end{frame}

\end{document}
"""
    
    file_path.write_text(latex_content)
    return file_path


@pytest.fixture
def large_file(temp_upload_dir: Path) -> Path:
    """Create large file for upload testing."""
    file_path = temp_upload_dir / "large_file.pdf"
    
    # Create a file larger than typical limit (15MB)
    size = 15 * 1024 * 1024
    with open(file_path, 'wb') as f:
        f.write(b'%PDF-1.4\n')
        f.seek(size - 1)
        f.write(b'\0')
    
    return file_path


@pytest.fixture
def malicious_file(temp_upload_dir: Path) -> Path:
    """Create file with malicious content for security testing."""
    file_path = temp_upload_dir / "malicious.pdf"
    
    # EICAR test virus signature
    eicar = b'X5O!P%@AP[4\\PZX54(P^)7CC)7}$EICAR-STANDARD-ANTIVIRUS-TEST-FILE!$H+H*'
    
    file_path.write_bytes(eicar)
    return file_path


@pytest_asyncio.fixture
async def websocket_client(client: AsyncClient):
    """Create WebSocket test client."""
    from websockets import connect
    
    # Get the test server URL
    base_url = str(client.base_url).replace("http://", "ws://")
    ws_url = f"{base_url}/api/v1/ws"
    
    async with connect(ws_url) as websocket:
        yield websocket


@pytest.fixture
def mock_ai_responses(mocker):
    """Mock AI service responses for testing."""
    # Mock Anthropic
    anthropic_mock = mocker.patch("app.services.ai.anthropic_provider.AnthropicProvider.generate")
    anthropic_mock.return_value = {
        "title": "Test Presentation",
        "slides": [
            {
                "type": "title",
                "title": "Test Presentation",
                "subtitle": "Integration Testing",
                "content": {"author": "Test Suite", "date": "2024"},
            },
            {
                "type": "content",
                "title": "Introduction",
                "content": {"text": "This is a test slide."},
            },
        ],
    }
    
    # Mock OpenAI
    openai_mock = mocker.patch("app.services.ai.openai_provider.OpenAIProvider.generate")
    openai_mock.return_value = {
        "title": "Fallback Presentation",
        "slides": [
            {
                "type": "title",
                "title": "Fallback Test",
                "content": {"text": "OpenAI fallback response"},
            },
        ],
    }
    
    return {
        "anthropic": anthropic_mock,
        "openai": openai_mock,
    }


@pytest.fixture
def test_presentation_data():
    """Sample presentation data for testing."""
    return {
        "title": "Machine Learning in Healthcare",
        "description": "A comprehensive overview of ML applications in healthcare",
        "audience": "Graduate students and researchers",
        "duration_minutes": 30,
        "slide_count": 15,
        "tone": "academic",
        "color_scheme": "professional",
        "include_citations": True,
        "sections": [
            {"title": "Introduction", "duration": 5},
            {"title": "Literature Review", "duration": 10},
            {"title": "Methodology", "duration": 8},
            {"title": "Results", "duration": 5},
            {"title": "Conclusion", "duration": 2},
        ],
    }


@pytest.fixture
def test_user_data():
    """Sample user data for testing."""
    return {
        "email": "test.user@university.edu",
        "password": "SecureTestP@ssw0rd!",
        "full_name": "Test User",
        "institution": "Test University",
        "role": "student",
    }


@pytest_asyncio.fixture
async def cleanup_database(db_session: AsyncSession):
    """Clean up database after tests."""
    yield
    
    # Clean up all test data
    await db_session.execute(text("TRUNCATE TABLE users CASCADE"))
    await db_session.execute(text("TRUNCATE TABLE presentations CASCADE"))
    await db_session.execute(text("TRUNCATE TABLE templates CASCADE"))
    await db_session.execute(text("TRUNCATE TABLE generation_jobs CASCADE"))
    await db_session.commit()